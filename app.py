from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from time import sleep
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
import schedule
import aspose.words as aw




class DollarQuote:
    
    #inicializa o driver,obtem a data atual e cria o nome do arquivo que será a imagem da página html
    
    def __init__(self):

        self.date = datetime.now().strftime("%d/%m/%Y")

        self.file_name = f'{self.date.replace("/","-")}.png'

        print(f'Inicializando bot hoje dia {self.date}...')
        
        self.chrome_options = Options()

        self.arguments = ['--lang=pt-BR', '--window-size=800,800', '--incognito']

        for argument in self.arguments:
            self.chrome_options.add_argument(argument)

        self.chrome_options.add_experimental_option("prefs", {
            'download.directory_upgrade': True,
            'download.prompt_for_download': False,
            "profile.default_content_setting_values.notifications": 2, 
            "profile.default_content_setting_values.automatic_downloads": 1,
        })

        self.driver = webdriver.Chrome(options=self.chrome_options)

        
    
    #obtem a cotaão do dolar
        
    def get_dollar_quote(self):
        print('Acessando site ...')
        sleep(1)
        self.driver.get("https://www.google.com/search?q=dolar")
        print('Carregando site ...')
        sleep(5)
        print('Obtendo cotação ...')
        self.cotacao_dolar = self.driver.find_element(By.CLASS_NAME,'DFlfde.SwHCTb').text.replace(',','.')
        print('Cotação obtida con sucesso!')
        return self.cotacao_dolar
        
    #salva uma imagem da página html  
    def save_screenshot(self):
        print('Salvando imagem do site ...')
        sleep(2)
        self.driver.save_screenshot(self.file_name)
        sleep(1)
        print('Imagem salva!')
        self.driver.close()
        self.driver.quit()
        
        


class DocxToPdfConverter:

    
    def __init__(self, texts, image_path, docx_filename, pdf_filename):
        # cria o arquivo docx e precisa dos parâmetros textopara o arquivo, caminho da imagem da página html,
        #nome do docx e nome do pdf que serão criados
        self.texts = texts
        self.image_path = image_path
        self.docx_filename = docx_filename
        self.pdf_filename = pdf_filename

    def write_docx(self,doc,texts):
        # escreve o arquivo docx
        print(f'Criando arquivo {self.docx_filename} ..')

        #texts é um dicionário
        for k,text in self.texts.items():
            # formatação para o título do texto
            if k == 'title':
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(text)
                run.bold = True
                run.font.size = Pt(20)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Depois de adicionar o 3º parágrafo, preciso adicionar a imagem da página html, pois logo
            #depois da imagem tenho que adicionar o último parágrafo
            elif k == 'paragraph3':
                doc.add_paragraph(text)
                paragraph = doc.add_paragraph()
                run = paragraph.add_run()
                run.add_picture(self.image_path, width=Inches(4.0))
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                
            else:
                doc.add_paragraph(text)
                
        #salva o arquivo docx
        
        doc.save(self.docx_filename)
        print('Arquivo criado!')
        
                          
    
    def create_docx(self):
        
        # cria o arquivo instanciando a classe Document
        doc = Document()
        # escreve o arquivo docx
        self.write_docx(doc,self.texts)
        
    def convert_docx_to_pdf(self,docx_path,name_pdf):
        
        # obtém o docx e faz a conversão para pdf
        print(f'Convetendo arquivo {docx_path} para {name_pdf}')
        doc = aw.Document(docx_path)
        #slava o arquivo pdf
        doc.save(name_pdf)
        print('Conversão realizada com sucesso!')
        
    # chama as funçoes anteriores para realizae os processos de criação e conversão
    def create_and_convert(self):

        self.create_docx()
        

        self.convert_docx_to_pdf(self.docx_filename,self.pdf_filename)
def main():
    # instanciando a classe DollarQuote
    bot =  DollarQuote()

    # obtendo cotação do dolar
    dollar_quote = bot.get_dollar_quote()
    # salvando o imagem do site
    bot.save_screenshot()

    # dicionário com os textos do docx
    texts = {
        'title': f'Cotação Atual do Dólar – {dollar_quote} ({bot.date})',
        'paragraph1':f'O dólar está no valor de {dollar_quote}, na data {bot.date}.',
        'paragraph2':'Valor cotado no site https://www.google.com/search?q=dolar',
        'paragraph3':'Print da cotação atual',
        'paragraph4':'Cotação feita por – Rafael Ribeiro'
        }

    #caminho(nome) da imagem do site
    image_path = bot.file_name
    #aproveitando o nome do arquivo png para nomear o docx
    docx_filename = bot.file_name.replace('.png','.docx')
    #aproveitando o nome do arquivo png para nomear o pdf
    pdf_filename = bot.file_name.replace('.png','.pdf')
    #instanciando a classe DocxToPdfConverter
    converter = DocxToPdfConverter(texts, image_path, docx_filename, pdf_filename)
    # criando o docx e convertendo para pdf
    converter.create_and_convert()
    print('Agendamento programado para amanhã às 17:00\n\n')

main()

schedule.every().day.at("17:00").do(main)

while True:
    schedule.run_pending()
    sleep(1)
